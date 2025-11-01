"""
Process ZBD (ЖБД) - Обробка CSV файлів та створення Word документів
Структура CSV файлу:
- Колонки 1-3: №, Звання, П.І.Б
- Колонки 4-34: Дні місяця (1-31) з описом подій

Створює Word документ з 31 таблицею (по одній для кожного дня місяця):
- Кожна таблиця: 3 колонки × мінімум 5 рядків
- Рядок 1: Заголовки ("Дата, час", "Завдання військ...", "Примітка")
- Рядок 2: Нумерація колонок (1, 2, 3)
- Рядок 3: Дата | Стандартний текст про стан підрозділу | Примітка
- Рядки 4+: Час події | Опис+список людей | Примітка (кожна подія - окремий рядок)
- Останній рядок: Підпис командира
"""

import sys
import json
import io
import csv
from pathlib import Path
from typing import List, Dict, Any, Set, Tuple, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from openpyxl import load_workbook

# Встановлюємо UTF-8 для stdout/stderr на Windows
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')


class ZBDProcessor:
    """Процесор для створення Word документів ЖБД з CSV файлів"""

    def __init__(self):
        self.csv_data = []
        self.month = None
        self.year = None
        self.config_excel_path = None
        self.commander_text = 'Командир РВП СпП 2 бСпП\nстарший лейтенант                                             Артем БУЛАВІН'
        self.position_headers = {}  # {позиція: текст для шапки}
        self.delivery_headers = {}  # {(майно, н.п.): текст для шапки}
        self.recon_headers = {}  # {вид пошуку: шаблон шапки з плейсхолдером}
        self.zvp_rotation_headers = {}  # {(повне_значення_пункту_4, блок): текст для шапки} - для спеціальних ротацій
        self.month_from_config = None
        self.daily_status_text = 'РВП СпП 2 БСпП продовжує виконання бойових завдань.\nСтан та положення – без змін'

    def read_csv(self, csv_path: str) -> List[List[str]]:
        """
        Читання CSV файлу

        Args:
            csv_path: Шлях до CSV файлу

        Returns:
            Список рядків CSV файлу
        """
        print(f"📖 Читання CSV файлу: {csv_path}")

        try:
            with open(csv_path, 'r', encoding='utf-8-sig') as f:
                reader = csv.reader(f)
                data = list(reader)

            print(f"✅ Прочитано {len(data)} рядків")
            
            # Зберігаємо дані в екземплярі класу
            self.csv_data = data
            
            # Визначаємо місяць та рік з назви файлу
            self._extract_month_year_from_filename(csv_path)
            
            return data

        except Exception as e:
            print(f"❌ Помилка читання CSV: {str(e)}", file=sys.stderr)
            raise

    def _extract_month_year_from_filename(self, csv_path: str) -> None:
        """
        Визначає місяць та рік з назви файлу або використовує поточну дату

        Args:
            csv_path: Шлях до CSV файлу
        """
        filename = Path(csv_path).stem
        
        # Словник українських місяців
        months_uk = {
            'січень': 1, 'січня': 1,
            'лютий': 2, 'лютого': 2,
            'березень': 3, 'березня': 3,
            'квітень': 4, 'квітня': 4,
            'травень': 5, 'травня': 5,
            'червень': 6, 'червня': 6,
            'липень': 7, 'липня': 7,
            'серпень': 8, 'серпня': 8,
            'вересень': 9, 'вересня': 9,
            'жовтень': 10, 'жовтня': 10,
            'листопад': 11, 'листопада': 11,
            'грудень': 12, 'грудня': 12
        }
        
        # Пробуємо знайти місяць в назві файлу
        filename_lower = filename.lower()
        for month_name, month_num in months_uk.items():
            if month_name in filename_lower:
                self.month = month_num
                break
        
        # Пробуємо знайти рік в назві файлу (4 цифри)
        import re
        year_match = re.search(r'20\d{2}', filename)
        if year_match:
            self.year = int(year_match.group())
        
        # Якщо не знайшли - використовуємо поточну дату
        if self.month is None or self.year is None:
            now = datetime.now()
            if self.month is None:
                self.month = now.month
            if self.year is None:
                self.year = now.year
        
        print(f"📅 Визначено місяць: {self.month}, рік: {self.year}")

    def read_config_excel(self, excel_path: str) -> None:
        """
        Читає конфігураційний Excel файл з листом "РВП"

        Args:
            excel_path: Шлях до Excel файлу
        """
        print(f"📖 Читання конфігураційного Excel: {excel_path}")
        
        try:
            self.config_excel_path = excel_path
            wb = load_workbook(excel_path, data_only=True)
            
            # Перевіряємо наявність листа "РВП"
            if 'РВП' not in wb.sheetnames:
                print(f"⚠️ Лист 'РВП' не знайдено в Excel файлі. Використовуються стандартні значення.")
                return
            
            ws = wb['РВП']
            
            # Читаємо A2 - текст командира
            commander_cell = ws['A2'].value
            if commander_cell:
                self.commander_text = str(commander_cell).strip()
                print(f"✅ Командир: {self.commander_text[:50]}...")
            
            # Читаємо E2 - текст щоденного стану
            daily_status_cell = ws['E2'].value
            if daily_status_cell:
                self.daily_status_text = str(daily_status_cell).strip()
                print(f"✅ Щоденний статус: {self.daily_status_text[:50]}...")
            
            # Читаємо маппінг позицій (колонки B, C)
            # Починаємо з рядка 2 (після заголовків)
            row = 2
            while True:
                position = ws.cell(row=row, column=2).value  # Колонка B
                header_text = ws.cell(row=row, column=3).value  # Колонка C
                month_value = ws.cell(row=row, column=4).value  # Колонка D
                
                # Маппінг доставок (колонки F, G, H)
                delivery_item = ws.cell(row=row, column=6).value  # Колонка F - майно
                delivery_location = ws.cell(row=row, column=7).value  # Колонка G - н.п.
                delivery_header = ws.cell(row=row, column=8).value  # Колонка H - текст шапки
                
                # Маппінг рекогностування (колонки I, J)
                recon_type = ws.cell(row=row, column=9).value  # Колонка I - вид пошуку
                recon_header = ws.cell(row=row, column=10).value  # Колонка J - шаблон шапки
                
                # Маппінг ротації для спеціальних позицій (колонки K, L, M)
                zvp_position = ws.cell(row=row, column=11).value  # Колонка K - повне значення пункту 4 (напр. "ЗВП Маріо-1")
                zvp_block = ws.cell(row=row, column=12).value  # Колонка L - блок (напр. "прибули:")
                zvp_header = ws.cell(row=row, column=13).value  # Колонка M - текст шапки
                
                # Якщо всі комірки пусті - закінчуємо
                if not position and not header_text and not delivery_item and not delivery_location and not recon_type and not zvp_position:
                    break
                
                # Зберігаємо маппінг позицій
                if position and header_text:
                    self.position_headers[str(position).strip()] = str(header_text).strip()
                    print(f"  📍 {position} → {header_text}")
                
                # Зберігаємо маппінг доставок
                if delivery_item and delivery_location and delivery_header:
                    key = (str(delivery_item).strip(), str(delivery_location).strip())
                    self.delivery_headers[key] = str(delivery_header).strip()
                    print(f"  📦 {delivery_item} → {delivery_location} → {delivery_header}")
                
                # Зберігаємо маппінг рекогностування
                if recon_type and recon_header:
                    # Зберігаємо як вид пошуку, так і блоки прибули/вибули
                    key = str(recon_type).strip().lower()
                    self.recon_headers[key] = str(recon_header).strip()
                    print(f"  🔍 {key} → {recon_header}")
                
                # Зберігаємо маппінг ротації для спеціальних позицій
                if zvp_position and zvp_block and zvp_header:
                    key = (str(zvp_position).strip(), str(zvp_block).strip().lower())
                    self.zvp_rotation_headers[key] = str(zvp_header).strip()
                    print(f"  🔄 {zvp_position} + {zvp_block} → {zvp_header}")
                
                # Зберігаємо місяць з першого рядка (якщо вказаний)
                if month_value and self.month_from_config is None:
                    try:
                        self.month_from_config = int(month_value)
                        print(f"📅 Місяць з конфігу: {self.month_from_config}")
                    except (ValueError, TypeError):
                        pass
                
                row += 1
            
            wb.close()
            print(f"✅ Конфігурація завантажена: {len(self.position_headers)} позицій")
            
        except Exception as e:
            print(f"⚠️ Помилка читання конфігураційного Excel: {str(e)}")
            print(f"⚠️ Використовуються стандартні значення")

    def _calculate_table_rows(self, events: List[Dict[str, Any]]) -> int:
        """
        Розраховує кількість рядків для таблиці

        Args:
            events: Список подій

        Returns:
            Кількість рядків
        """
        if not events:
            return 5  # Мінімум 5 рядків: заголовок, нумерація, дата, порожній, підпис
        
        # 2 заголовки + 1 рядок стандартного тексту + події + 1 підпис
        # Кожна подія займає 1 рядок в таблиці
        return 3 + len(events) + 1

    def _fill_events(self, table, events: List[Dict[str, Any]], start_row: int) -> int:
        """
        Заповнює таблицю подіями (починаючи з рядка 4)

        Args:
            table: Таблиця Word
            events: Список подій
            start_row: Початковий рядок (індекс рядка 4)

        Returns:
            Номер наступного вільного рядка
        """
        current_row = start_row
        
        print(f"    📝 Заповнення {len(events)} подій, початковий рядок: {start_row}")
        
        for event in events:
            event_type = event.get('type', 'обстріл')
            print(f"      → Заповнення події '{event_type}' о {event.get('time')} в рядок {current_row}")
            
            if event_type == 'обстріл':
                current_row = self._fill_obstril_event(table, event, current_row)
            elif event_type == 'ротація':
                current_row = self._fill_rotation_event(table, event, current_row)
            elif event_type == 'доставка':
                current_row = self._fill_delivery_event(table, event, current_row)
            elif event_type == 'рекогностування':
                current_row = self._fill_recon_event(table, event, current_row)
        
        print(f"    ✅ Заповнення завершено, наступний рядок: {current_row}")
        return current_row

    def _fill_obstril_event(self, table, event: Dict[str, Any], row: int) -> int:
        """
        Заповнює рядок таблиці подією типу "обстріл"

        Args:
            table: Таблиця Word
            event: Інформація про подію
            row: Номер рядка

        Returns:
            Номер наступного вільного рядка
        """
        # Формуємо текст для колонки 1 (час)
        col1_text = event['time']
        
        # Формуємо текст для колонки 2 (опис + список людей)
        col2_lines = [event['description']]
        
        # Якщо це обстріл РВП і є люди
        if event.get('is_rvp') and event.get('people'):
            col2_lines.append('На РЗ РВП перебували:')
            # Додаємо список людей
            for person in event['people']:
                col2_lines.append(person)
        
        # Записуємо в комірки
        cell_col1 = table.rows[row].cells[0]
        cell_col1.text = col1_text
        self._format_cell_text(cell_col1, col1_text)
        
        cell_col2 = table.rows[row].cells[1]
        cell_col2.text = '\n'.join(col2_lines)
        self._format_cell_text(cell_col2, '\n'.join(col2_lines))
        
        return row + 1

    def _fill_rotation_event(self, table, event: Dict[str, Any], row: int) -> int:
        """
        Заповнює рядок таблиці подією типу "ротація"

        Args:
            table: Таблиця Word
            event: Інформація про подію
            row: Номер рядка

        Returns:
            Номер наступного вільного рядка
        """
        # Формуємо текст для колонки 1 (час)
        col1_text = event['time']
        
        # Формуємо текст для колонки 2
        col2_lines = []
        
        position = event.get('position', '')
        
        # Перевіряємо чи є спеціальні шапки для цієї позиції
        has_special_header = False
        if position:
            # Перевіряємо чи є хоча б один блок зі спеціальною шапкою
            # Для вибули - тільки якщо немає прибули
            if (position, 'прибули:') in self.zvp_rotation_headers:
                has_special_header = True
            elif not event.get('arrived') and (position, 'вибули:') in self.zvp_rotation_headers:
                has_special_header = True
        
        # Шапка з позицією - виводимо ТІЛЬКИ якщо немає спеціальних шапок
        if position and not has_special_header:
            # Перевіряємо чи є кастомний текст для цієї позиції в конфігу
            if position in self.position_headers:
                header_text = self.position_headers[position]
            else:
                header_text = f"Проведено ротацію о/с ВП \"{position}\""
            col2_lines.append(header_text)
        
        # Прибули - перевіряємо спеціальну шапку
        if event.get('arrived'):
            zvp_key = (position, 'прибули:')
            
            if zvp_key in self.zvp_rotation_headers:
                # Використовуємо спеціальну шапку з конфігу
                col2_lines.append(self.zvp_rotation_headers[zvp_key])
                col2_lines.append('прибули:')
            else:
                col2_lines.append('прибули:')
            
            for person in event['arrived']:
                col2_lines.append(person)
        
        # Вибули - перевіряємо спеціальну шапку (тільки якщо немає прибули)
        if event.get('departed'):
            zvp_key = (position, 'вибули:')
            
            if zvp_key in self.zvp_rotation_headers and not event.get('arrived'):
                # Використовуємо спеціальну шапку з конфігу (тільки якщо немає прибули)
                col2_lines.append(self.zvp_rotation_headers[zvp_key])
                col2_lines.append('вибули:')
            else:
                col2_lines.append('вибули:')
            
            for person in event['departed']:
                col2_lines.append(person)
        
        # Водії
        if event.get('drivers'):
            col2_lines.append('водій:')
            for driver in event['drivers']:
                col2_lines.append(driver)
        
        # Штурмани
        if event.get('navigators'):
            col2_lines.append('штурман:')
            for navigator in event['navigators']:
                col2_lines.append(navigator)
        
        # Записуємо в комірки
        cell_col1 = table.rows[row].cells[0]
        cell_col1.text = col1_text
        self._format_cell_text(cell_col1, col1_text)
        
        cell_col2 = table.rows[row].cells[1]
        cell_col2.text = '\n'.join(col2_lines)
        self._format_cell_text(cell_col2, '\n'.join(col2_lines))
        
        return row + 1

    def _fill_delivery_event(self, table, event: Dict[str, Any], row: int) -> int:
        """
        Заповнює рядок таблиці подією типу "доставка"

        Args:
            table: Таблиця Word
            event: Інформація про подію
            row: Номер рядка

        Returns:
            Номер наступного вільного рядка
        """
        # Формуємо текст для колонки 1 (час)
        col1_text = event['time']
        
        # Формуємо текст для колонки 2
        col2_lines = []
        
        # Шапка - перевіряємо чи є кастомний текст для цього майна та н.п.
        cargo = event.get('cargo')
        location = event.get('location')
        
        if cargo and location:
            key = (cargo, location)
            if key in self.delivery_headers:
                header_text = self.delivery_headers[key]
            else:
                header_text = f"Здійснено доставку: {cargo}, н.п. {location}"
            col2_lines.append(header_text)
        
        # Загальний список (якщо є) - спочатку
        if event.get('general'):
            for person in event['general']:
                col2_lines.append(person)
        
        # Водії
        if event.get('drivers'):
            col2_lines.append('водій:')
            for driver in event['drivers']:
                col2_lines.append(driver)
        
        # Штурмани - в кінці
        if event.get('navigators'):
            col2_lines.append('штурман:')
            for navigator in event['navigators']:
                col2_lines.append(navigator)
        
        # Записуємо в комірки
        cell_col1 = table.rows[row].cells[0]
        cell_col1.text = col1_text
        self._format_cell_text(cell_col1, col1_text)
        
        cell_col2 = table.rows[row].cells[1]
        cell_col2.text = '\n'.join(col2_lines)
        self._format_cell_text(cell_col2, '\n'.join(col2_lines))
        
        return row + 1

    def _fill_recon_event(self, table, event: Dict[str, Any], row: int) -> int:
        """
        Заповнює рядок таблиці подією типу "рекогностування"

        Args:
            table: Таблиця Word
            event: Інформація про подію
            row: Номер рядка

        Returns:
            Номер наступного вільного рядка
        """
        # Формуємо текст для колонки 1 (час)
        col1_text = event['time']
        
        # Формуємо текст для колонки 2
        col2_lines = []
        
        # Шапка
        recon_type = event.get('recon_type')
        location = event.get('location', '')
        
        # Якщо є вид пошуку (напр. "поточно!") - перевіряємо конфіг
        if recon_type:
            recon_type_key = recon_type.lower()
            if recon_type_key in self.recon_headers:
                # Використовуємо шаблон з конфігу і замінюємо плейсхолдер на н.п.
                header_template = self.recon_headers[recon_type_key]
                header_text = header_template.replace('(значення з пункту 4)', location)
                header_text = header_text.replace('(значення з рункту 4)', location)
                col2_lines.append(header_text)
            else:
                # Якщо немає в конфігу - стандартний текст
                col2_lines.append(f"Рекогностування ({recon_type}), н.п. {location}")
            
            # Додаємо блоки прибули/вибули НИЖЧЕ шапки якщо є
            if event.get('arrived'):
                # Шукаємо заголовок для блоку "прибули:" в конфігу
                if 'прибули:' in self.recon_headers:
                    arrived_header = self.recon_headers['прибули:']
                    arrived_header = arrived_header.replace('(значення з пункту 4)', location)
                    arrived_header = arrived_header.replace('(значення з рункту 4)', location)
                    col2_lines.append(arrived_header)
                    col2_lines.append('прибули:')
                else:
                    col2_lines.append('прибули:')
                for person in event['arrived']:
                    col2_lines.append(person)
            
            if event.get('departed'):
                # Шукаємо заголовок для блоку "вибули:" в конфігу
                if 'вибули:' in self.recon_headers:
                    departed_header = self.recon_headers['вибули:']
                    departed_header = departed_header.replace('(значення з пункту 4)', location)
                    departed_header = departed_header.replace('(значення з рункту 4)', location)
                    col2_lines.append(departed_header)
                    col2_lines.append('вибули:')
                else:
                    col2_lines.append('вибули:')
                for person in event['departed']:
                    col2_lines.append(person)
        else:
            # Якщо це просто прибули/вибули без виду пошуку
            
            if event.get('arrived'):
                # Шукаємо заголовок для блоку "прибули:" в конфігу
                if 'прибули:' in self.recon_headers:
                    arrived_header = self.recon_headers['прибули:']
                    arrived_header = arrived_header.replace('(значення з пункту 4)', location)
                    arrived_header = arrived_header.replace('(значення з рункту 4)', location)
                    col2_lines.append(arrived_header)
                    col2_lines.append('прибули:')
                else:
                    col2_lines.append(f"Рекогностування н.п. {location}")
                    col2_lines.append('прибули:')
                for person in event['arrived']:
                    col2_lines.append(person)
            
            if event.get('departed'):
                # Шукаємо заголовок для блоку "вибули:" в конфігу
                if 'вибули:' in self.recon_headers:
                    departed_header = self.recon_headers['вибули:']
                    departed_header = departed_header.replace('(значення з пункту 4)', location)
                    departed_header = departed_header.replace('(значення з рункту 4)', location)
                    col2_lines.append(departed_header)
                    col2_lines.append('вибули:')
                else:
                    if not event.get('arrived'):
                        col2_lines.append(f"Рекогностування н.п. {location}")
                    col2_lines.append('вибули:')
                for person in event['departed']:
                    col2_lines.append(person)
        
        # Записуємо в комірки
        cell_col1 = table.rows[row].cells[0]
        cell_col1.text = col1_text
        self._format_cell_text(cell_col1, col1_text)
        
        cell_col2 = table.rows[row].cells[1]
        cell_col2.text = '\n'.join(col2_lines)
        self._format_cell_text(cell_col2, '\n'.join(col2_lines))
        
        return row + 1

    def _format_cell_text(self, cell, text: str) -> None:
        """
        Форматує текст в комірці

        Args:
            cell: Комірка таблиці
            text: Текст для форматування
        """
        if not cell.paragraphs:
            return
        
        paragraph = cell.paragraphs[0]
        
        # Очищаємо параграф
        paragraph.clear()
        
        # Додаємо текст з форматуванням
        run = paragraph.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    def _get_events_for_day(self, data: List[List[str]], day: int) -> List[Dict[str, Any]]:
        """
        Витягує всі події для конкретного дня з розбором структури

        Args:
            data: Дані CSV файлу
            day: Номер дня місяця (1-31)

        Returns:
            Список подій з розбором часу, інформації та списку людей
        """
        day_col_idx = 3 + (day - 1)  # Колонка для цього дня
        events_dict = {}  # Словник для групування подій
        
        print(f"  🔍 Аналіз дня {day}, колонка індекс {day_col_idx}")
        
        # Пропускаємо перший рядок (заголовки)
        for row_idx, row in enumerate(data[1:], start=1):
            # Перевіряємо, чи є ця колонка в рядку
            if day_col_idx >= len(row):
                continue
                
            cell_value = row[day_col_idx].strip()
            
            if not cell_value or len(cell_value) <= 2:
                continue
            
            print(f"    📄 Рядок {row_idx}: знайдено текст довжиною {len(cell_value)} символів")
            
            print(f"    📄 Рядок {row_idx}: знайдено текст довжиною {len(cell_value)} символів")
            
            # Отримуємо звання та ПІБ з колонок 1 та 2
            rank = row[1].strip() if len(row) > 1 else ''
            full_name = row[2].strip() if len(row) > 2 else ''
            person = f"{rank} {full_name}".strip()
            
            # Парсимо інформацію про подію
            event_info = self._parse_event_info(cell_value)
            
            if event_info:
                event_type = event_info.get('type', 'обстріл')
                
                if event_type == 'обстріл':
                    # Обстріли: створюємо унікальний ключ для групування однакових
                    event_key = f"obstril_{event_info['time']}_{event_info.get('description', '')}"
                    
                    if event_key not in events_dict:
                        events_dict[event_key] = {
                            'type': 'обстріл',
                            'time': event_info['time'],
                            'description': event_info.get('description', ''),
                            'is_rvp': event_info.get('is_rvp', False),
                            'people': [],
                            'is_separate': True
                        }
                    
                    # Додаємо людину до списку
                    if person:
                        events_dict[event_key]['people'].append(person)
                
                elif event_type == 'ротація':
                    # Ротації: групуємо за часом та позицією
                    event_key = f"rotation_{event_info['time']}_{event_info.get('position', '')}"
                    
                    if event_key not in events_dict:
                        events_dict[event_key] = {
                            'type': 'ротація',
                            'time': event_info['time'],
                            'position': event_info.get('position'),  # Повна позиція (напр. "ЗВП Маріо-1")
                            'arrived': [],
                            'departed': [],
                            'drivers': [],
                            'navigators': [],
                            'is_separate': False
                        }
                    
                    # Додаємо людину до відповідного списку (звання + ПІБ з CSV)
                    # Визначаємо блок по block_type з event_info
                    if person:  # person = звання + ПІБ з колонок 1 та 2
                        block_type = event_info.get('block_type')
                        
                        # Прибули
                        if block_type == 'arrived':
                            events_dict[event_key]['arrived'].append(person)
                        # Вибули
                        elif block_type == 'departed':
                            events_dict[event_key]['departed'].append(person)
                        # Водій
                        elif block_type == 'driver':
                            events_dict[event_key]['drivers'].append(person)
                        # Штурман
                        elif block_type == 'navigator':
                            events_dict[event_key]['navigators'].append(person)
                    
                    # Додаємо водіїв та штурманів з самого event_info (якщо вони вказані прямо в CSV, а не через людей)
                    # Це на випадок якщо в одній ячейці CSV буде кілька людей в блоці водій/штурман
                    if event_info.get('drivers') and not person:
                        events_dict[event_key]['drivers'].extend(event_info['drivers'])
                    if event_info.get('navigators') and not person:
                        events_dict[event_key]['navigators'].extend(event_info['navigators'])
                
                elif event_type == 'доставка':
                    # Доставки: групуємо за часом, майном та н.п.
                    event_key = f"delivery_{event_info['time']}_{event_info.get('cargo', '')}_{event_info.get('location', '')}"
                    
                    if event_key not in events_dict:
                        events_dict[event_key] = {
                            'type': 'доставка',
                            'time': event_info['time'],
                            'cargo': event_info.get('cargo'),
                            'location': event_info.get('location'),
                            'drivers': [],
                            'navigators': [],
                            'general': [],  # Загальний список
                            'is_separate': False
                        }
                    
                    # Додаємо людину до відповідного списку (звання + ПІБ з CSV)
                    if person:
                        block_type = event_info.get('block_type')
                        
                        # Водій
                        if block_type == 'driver':
                            events_dict[event_key]['drivers'].append(person)
                        # Штурман
                        elif block_type == 'navigator':
                            events_dict[event_key]['navigators'].append(person)
                        # Загальний список (якщо блок не визначений або порожній)
                        else:
                            events_dict[event_key]['general'].append(person)
                
                elif event_type == 'рекогностування':
                    # Рекогностування: групуємо за часом, типом пошуку та н.п.
                    recon_type = event_info.get('recon_type', '')
                    location = event_info.get('location', '')
                    event_key = f"recon_{event_info['time']}_{recon_type}_{location}"
                    
                    if event_key not in events_dict:
                        events_dict[event_key] = {
                            'type': 'рекогностування',
                            'time': event_info['time'],
                            'recon_type': recon_type,
                            'location': location,
                            'arrived': [],
                            'departed': [],
                            'is_separate': False
                        }
                    
                    # Додаємо людину до відповідного списку
                    if person:
                        block_type = event_info.get('block_type')
                        
                        if block_type == 'arrived':
                            events_dict[event_key]['arrived'].append(person)
                        elif block_type == 'departed':
                            events_dict[event_key]['departed'].append(person)
        
        # Повертаємо список подій, відсортованих за часом
        events_list = list(events_dict.values())
        events_list.sort(key=lambda x: x['time'])
        
        if events_list:
            print(f"  ✅ Знайдено {len(events_list)} подій для дня {day}")
        
        return events_list

    def _parse_event_info(self, text: str) -> Optional[Dict[str, Any]]:
        """
        Парсить інформацію про подію

        Args:
            text: Текст з інформацією про подію

        Returns:
            Словник з розбором події або None
        """
        lines = text.split('\n')
        
        if len(lines) < 2:
            return None
        
        # Визначаємо тип події
        # Формат може бути:
        # 1) "Обстріл\n1) час\nопис" (старий формат)
        # 2) "1) час\n2) опис" (новий формат - визначаємо тип по змісту)
        
        first_line = lines[0].strip()
        event_type = None
        
        # Перевіряємо старий формат (перший рядок - тип події)
        if 'обстріл' in first_line.lower():
            event_type = 'обстріл'
        elif 'ротац' in first_line.lower():
            event_type = 'ротація'
        elif 'рекогност' in first_line.lower():
            event_type = 'рекогностування'
        elif 'доставка' in first_line.lower():
            event_type = 'доставка'
        
        # Витягуємо час
        time = None
        time_line_idx = 0
        
        if event_type:
            # Старий формат: час у другому рядку
            time_line = lines[1].strip() if len(lines) > 1 else ''
            time_match = re.search(r'(\d{2}:\d{2}(?:-\d{2}:\d{2})?)', time_line)
            if time_match:
                time = time_match.group(1)
                time_line_idx = 1
        else:
            # Новий формат: час у першому рядку (формат "1) 23:08-23:14")
            time_match = re.search(r'(\d{2}:\d{2}(?:-\d{2}:\d{2})?)', first_line)
            if time_match:
                time = time_match.group(1)
                time_line_idx = 0
                
                # Визначаємо тип по змісту в наступних рядках
                full_text = '\n'.join(lines).lower()
                if 'обстріл' in full_text or 'по рз рвп' in full_text or 'по рвп' in full_text:
                    event_type = 'обстріл'
                elif 'ротац' in full_text:
                    event_type = 'ротація'
                elif 'рекогност' in full_text:
                    event_type = 'рекогностування'
                elif 'доставка' in full_text:
                    event_type = 'доставка'
        
        if not time or not event_type:
            return None
        
        if event_type == 'обстріл':
            return self._parse_obstril(lines, time, time_line_idx)
        elif event_type == 'ротація':
            return self._parse_rotation(lines, time, time_line_idx)
        elif event_type == 'рекогностування':
            return self._parse_recon(lines, time, time_line_idx)
        elif event_type == 'доставка':
            return self._parse_delivery(lines, time, time_line_idx)
        
        return None

    def _parse_obstril(self, lines: List[str], time: str, time_line_idx: int = 1) -> Optional[Dict[str, Any]]:
        """
        Парсить інформацію про обстріл

        Args:
            lines: Рядки тексту події
            time: Час події
            time_line_idx: Індекс рядка з часом (0 або 1)

        Returns:
            Словник з інформацією про обстріл
        """
        # Витягуємо опис (рядок після часу)
        description = ''
        
        # Починаємо шукати опис після рядка з часом
        for line in lines[time_line_idx + 1:]:
            # Видаляємо номер на початку рядка якщо є (формат "2) опис")
            clean_line = re.sub(r'^\d+\)\s*', '', line.strip())
            if clean_line:
                description = clean_line
                break
        
        if not description:
            return None
        
        # Перевіряємо, чи це обстріл РВП (шукаємо "по рз рвп" або "по рвп" в описі)
        is_rvp = ('по рз рвп' in description.lower() or 
                  'по рвп' in description.lower())
        
        return {
            'type': 'обстріл',
            'time': time,
            'description': description,
            'is_rvp': is_rvp,
            'is_separate': True  # Обстріл завжди окремо
        }

    def _parse_rotation(self, lines: List[str], time: str, time_line_idx: int = 1) -> Optional[Dict[str, Any]]:
        """
        Парсить інформацію про ротацію

        Args:
            lines: Рядки тексту події
            time: Час події
            time_line_idx: Індекс рядка з часом (0 або 1)

        Returns:
            Словник з інформацією про ротацію
        """
        # Шукаємо блоки: вибули, прибули, водій, штурман
        position = None
        block_type = None  # 'arrived', 'departed', 'driver', 'navigator'
        drivers = []
        navigators = []
        found_block = False  # Чи знайшли блок (для визначення позиції)
        
        # Починаємо після рядка з часом
        for line in lines[time_line_idx + 1:]:
            # Видаляємо номер на початку рядка
            clean_line = re.sub(r'^\d+\)\s*', '', line.strip())
            if not clean_line:
                continue
            
            line_lower = clean_line.lower()
            
            # Якщо вже знайшли блок і ще не має позиції - цей рядок є позицією
            if found_block and not position:
                position = clean_line  # Зберігаємо повну позицію (напр. "ЗВП Маріо-1")
                found_block = False  # Скидаємо прапорець
                continue
            
            # Визначаємо тип блоку
            if 'прибули:' in line_lower:
                block_type = 'arrived'
                found_block = True
                # Перевіряємо, чи є позиція в тому ж рядку після двокрапки
                parts = clean_line.split(':', 1)
                if len(parts) > 1 and parts[1].strip():
                    position = parts[1].strip()  # Повна позиція
                    found_block = False
            elif 'вибули:' in line_lower:
                block_type = 'departed'
                found_block = True
                # Перевіряємо, чи є позиція в тому ж рядку після двокрапки
                parts = clean_line.split(':', 1)
                if len(parts) > 1 and parts[1].strip():
                    position = parts[1].strip()  # Повна позиція
                    found_block = False
            elif 'водій:' in line_lower:
                block_type = 'driver'
                found_block = True
                # Перевіряємо, чи є дані в тому ж рядку після двокрапки
                parts = clean_line.split(':', 1)
                if len(parts) > 1 and parts[1].strip():
                    drivers.append(parts[1].strip())
                    found_block = False
            elif 'штурман:' in line_lower:
                block_type = 'navigator'
                found_block = True
                # Перевіряємо, чи є дані в тому ж рядку після двокрапки
                parts = clean_line.split(':', 1)
                if len(parts) > 1 and parts[1].strip():
                    navigators.append(parts[1].strip())
                    found_block = False
            elif position and block_type in ['driver', 'navigator']:
                # Додаємо людей до відповідного блоку (тільки після того як знайшли позицію)
                if block_type == 'driver':
                    drivers.append(clean_line)
                elif block_type == 'navigator':
                    navigators.append(clean_line)
        
        return {
            'type': 'ротація',
            'time': time,
            'position': position,  # Повна позиція для пошуку в конфігу (напр. "ЗВП Маріо-1")
            'block_type': block_type,  # Тепер може бути 'arrived', 'departed', 'driver', 'navigator'
            'drivers': drivers,
            'navigators': navigators,
            'is_separate': False  # Може групуватися з іншими за часом
        }

    def _parse_recon(self, lines: List[str], time: str, time_line_idx: int = 1) -> Optional[Dict[str, Any]]:
        """
        Парсить інформацію про рекогностування

        Args:
            lines: Рядки тексту події
            time: Час події
            time_line_idx: Індекс рядка з часом (0 або 1)

        Returns:
            Словник з інформацією про рекогностування
        """
        # Структура:
        # Варіант 1: Рядок після часу: "прибули:" → наступний рядок: н.п.
        # Варіант 2: Рядок після часу: "вибули:" → наступний рядок: н.п.
        # Варіант 3: Рядок після часу: вид пошуку (напр. "поточно!") → наступний рядок: н.п.
        
        block_type = None  # 'arrived', 'departed', або None
        recon_type = None  # Вид пошуку (напр. "поточно!")
        location = None  # Н.п.
        found_block = False
        
        # Починаємо після рядка з часом
        for line in lines[time_line_idx + 1:]:
            clean_line = re.sub(r'^\d+\)\s*', '', line.strip())
            if not clean_line:
                continue
            
            line_lower = clean_line.lower()
            
            # Якщо вже знайшли блок/тип і ще немає н.п. - цей рядок є н.п.
            if found_block and not location:
                location = clean_line
                break
            
            # Визначаємо тип блоку
            if 'прибули:' in line_lower:
                block_type = 'arrived'
                found_block = True
            elif 'вибули:' in line_lower:
                block_type = 'departed'
                found_block = True
            else:
                # Це вид пошуку (напр. "поточно!")
                recon_type = clean_line
                found_block = True
        
        return {
            'type': 'рекогностування',
            'time': time,
            'block_type': block_type,  # 'arrived', 'departed', або None
            'recon_type': recon_type,  # Вид пошуку або None
            'location': location,  # Н.п.
            'is_separate': False
        }

    def _parse_delivery(self, lines: List[str], time: str, time_line_idx: int = 1) -> Optional[Dict[str, Any]]:
        """
        Парсить інформацію про доставку

        Args:
            lines: Рядки тексту події
            time: Час події
            time_line_idx: Індекс рядка з часом (0 або 1)

        Returns:
            Словник з інформацією про доставку
        """
        # Структура:
        # Рядок після часу: майно
        # Наступний рядок: н.п. (населений пункт)
        # Далі: блоки (водій:, штурман:, або пусто для загального списку)
        
        cargo = None  # Майно
        location = None  # Н.п.
        block_type = None  # 'driver', 'navigator', або None (загальний список)
        drivers = []
        navigators = []
        general = []  # Загальний список (ті хто не водій і не штурман)
        found_block = False
        
        # Починаємо після рядка з часом
        for idx, line in enumerate(lines[time_line_idx + 1:]):
            clean_line = re.sub(r'^\d+\)\s*', '', line.strip())
            if not clean_line:
                continue
            
            line_lower = clean_line.lower()
            
            # Перший рядок - майно
            if cargo is None:
                cargo = clean_line
                continue
            
            # Другий рядок - н.п.
            if location is None:
                location = clean_line
                continue
            
            # Якщо вже знайшли блок і він порожній - наступний рядок йде в загальний список
            if found_block and block_type is None:
                general.append(clean_line)
                continue
            
            # Визначаємо блок
            if 'водій:' in line_lower:
                block_type = 'driver'
                found_block = True
                parts = clean_line.split(':', 1)
                if len(parts) > 1 and parts[1].strip():
                    drivers.append(parts[1].strip())
                    found_block = False
            elif 'штурман:' in line_lower:
                block_type = 'navigator'
                found_block = True
                parts = clean_line.split(':', 1)
                if len(parts) > 1 and parts[1].strip():
                    navigators.append(parts[1].strip())
                    found_block = False
            elif block_type == 'driver':
                drivers.append(clean_line)
            elif block_type == 'navigator':
                navigators.append(clean_line)
            elif block_type is None:
                # Якщо блок не визначений і рядок не містить ":", додаємо в загальний список
                if ':' not in clean_line:
                    general.append(clean_line)
        
        return {
            'type': 'доставка',
            'time': time,
            'cargo': cargo,
            'location': location,
            'block_type': block_type,
            'drivers': drivers,
            'navigators': navigators,
            'general': general,
            'is_separate': False
        }

    def create_word_document(self, output_path: str) -> None:
        """
        Створення Word документу з таблицями для кожного дня

        Args:
            output_path: Шлях для збереження Word документу
        """
        print(f"\n📝 Створення Word документу...")

        try:
            # Створюємо новий документ
            doc = Document()

            # Налаштовуємо поля сторінки
            print("📐 Налаштування полів сторінки...")
            section = doc.sections[0]
            section.top_margin = Inches(1.5 / 2.54)      # 1.5 см
            section.bottom_margin = Inches(1.5 / 2.54)   # 1.5 см
            section.left_margin = Inches(2.5 / 2.54)     # 2.5 см
            section.right_margin = Inches(1.5 / 2.54)    # 1.5 см

            # Створюємо таблицю для кожного дня (1-31)
            print(f"📊 Створення таблиць для всіх днів місяця (1-31)...")
            
            # Використовуємо місяць з конфігу якщо є, інакше з CSV
            display_month = self.month_from_config if self.month_from_config else self.month
            
            for day in range(1, 32):
                print(f"\n📅 Створення таблиці для дня {day}.{display_month:02d}.{self.year}...")
                
                # Отримуємо події для цього дня
                events = self._get_events_for_day(self.csv_data, day)
                
                # Розраховуємо кількість рядків таблиці
                total_rows = self._calculate_table_rows(events)
                
                # Додаємо таблицю з динамічною кількістю рядків (мінімум 5)
                num_rows = max(5, total_rows)
                table = doc.add_table(rows=num_rows, cols=3)
                table.style = 'Table Grid'

                # Встановлюємо ширину колонок
                column_widths = [
                    2.48,   # Колонка 1: 2.48 см
                    12.32,  # Колонка 2: 12.32 см
                    1.99    # Колонка 3: 1.99 см
                ]

                for row in table.rows:
                    for col_idx, cell in enumerate(row.cells):
                        cell.width = Inches(column_widths[col_idx] / 2.54)

                # Рядок 1: Заголовки колонок
                headers = ['Дата,\nчас', 'Завдання військ та стисле висвітлення ходу бойових дій', 'Примітка']
                for col_idx, header in enumerate(headers):
                    cell = table.rows[0].cells[col_idx]
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.text = ''
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(header)
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'

                # Рядок 2: Нумерація колонок
                numbers = ['1', '2', '3']
                for col_idx, number in enumerate(numbers):
                    cell = table.rows[1].cells[col_idx]
                    cell.text = number
                    if cell.paragraphs:
                        paragraph = cell.paragraphs[0]
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(number)
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Рядок 3: Дата та стандартний текст (завжди присутній)
                date_str = f"{day:02d}.{display_month:02d}.{self.year}"
                row3_data = [
                    date_str,
                    self.daily_status_text,  # Використовуємо текст з конфігу або стандартний
                    ''  # Примітка порожня
                ]
                
                for col_idx, cell_data in enumerate(row3_data):
                    cell = table.rows[2].cells[col_idx]
                    cell.text = cell_data
                    self._format_cell_text(cell, cell_data)
                
                # Рядки 4+: Події (якщо є)
                current_row = 3  # Починаємо з рядка 4 (індекс 3)
                if events:
                    current_row = self._fill_events(table, events, current_row)

                # Останній рядок: Командир та підпис (завжди в передостанньому рядку)
                signature_row = num_rows - 1
                row_signature = [
                    '',  # Колонка 1 порожня
                    self.commander_text,  # Використовуємо текст з конфігу або стандартний
                    ''  # Колонка 3 порожня
                ]

                for col_idx, cell_data in enumerate(row_signature):
                    cell = table.rows[signature_row].cells[col_idx]
                    cell.text = cell_data
                    if cell.paragraphs:
                        paragraph = cell.paragraphs[0]
                        if paragraph.runs:
                            run = paragraph.runs[0]
                        else:
                            run = paragraph.add_run(cell_data)
                        run.font.bold = True  # Жирний шрифт для командира
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'

                # Додаємо кордони до таблиці
                self._set_table_borders(table)
                
                # Додаємо розрив сторінки після таблиці (крім останньої)
                if day < 31:
                    doc.add_page_break()

            # Зберігаємо документ
            print(f"\n💾 Збереження документу: {output_path}")
            doc.save(output_path)
            print("✅ Word документ успішно створено!")

        except Exception as e:
            print(f"❌ Помилка створення Word документу: {str(e)}", file=sys.stderr)
            raise

    def _set_table_borders(self, table):
        """
        Встановлює кордони для таблиці

        Args:
            table: Таблиця docx
        """
        tbl = table._element
        tblPr = tbl.tblPr

        # Створюємо елемент borders якщо його немає
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)

        # Встановлюємо всі кордони
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Товщина кордону
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Чорний колір
            tblBorders.append(border)

    def process(self, csv_path: str, output_path: str) -> Dict[str, Any]:
        """
        Головна функція обробки

        Args:
            csv_path: Шлях до CSV файлу
            output_path: Шлях для збереження Word документу

        Returns:
            Словник з результатами обробки
        """
        try:
            print("🚀 === ОБРОБКА ЖБД ===\n")

            # Читаємо CSV файл
            self.csv_data = self.read_csv(csv_path)

            # Створюємо Word документ з таблицями для всіх днів
            self.create_word_document(output_path)

            print("\n✅ === ОБРОБКА ЗАВЕРШЕНА ===")

            return {
                'success': True,
                'output_path': output_path,
                'rows_processed': 31,
                'message': f'Успішно створено Word документ з 31 таблицею (по одній на кожен день місяця)'
            }

        except Exception as e:
            error_msg = f"Помилка обробки: {str(e)}"
            print(f"\n❌ {error_msg}", file=sys.stderr)
            return {
                'success': False,
                'error': error_msg
            }


def main():
    """Головна функція - приймає JSON конфігурацію через stdin"""

    try:
        # Читаємо конфігурацію з stdin
        config_json = sys.stdin.read()
        config = json.loads(config_json)

        csv_path = config.get('csv_path')
        output_path = config.get('output_path')
        config_excel_path = config.get('config_excel_path')  # Додано: шлях до конфігураційного Excel

        # Валідація вхідних параметрів
        if not csv_path:
            raise ValueError("Не вказано шлях до CSV файлу")

        if not output_path:
            raise ValueError("Не вказано шлях для збереження результату")

        # Перевірка існування CSV файлу
        if not Path(csv_path).exists():
            raise FileNotFoundError(f"CSV файл не знайдено: {csv_path}")

        # Створюємо директорію для результату якщо не існує
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)

        # Обробка
        processor = ZBDProcessor()
        
        # Якщо вказаний конфігураційний Excel - читаємо його
        if config_excel_path and Path(config_excel_path).exists():
            processor.read_config_excel(config_excel_path)
        
        result = processor.process(csv_path, output_path)

        # Повертаємо результат у JSON форматі
        print(f"\n__RESULT__{json.dumps(result, ensure_ascii=False)}__END__")

        sys.exit(0 if result['success'] else 1)

    except Exception as e:
        error_msg = str(e)
        print(f"\n❌ КРИТИЧНА ПОМИЛКА: {error_msg}", file=sys.stderr)

        result = {
            'success': False,
            'error': error_msg
        }
        print(f"\n__RESULT__{json.dumps(result, ensure_ascii=False)}__END__")
        sys.exit(1)


if __name__ == '__main__':
    main()
